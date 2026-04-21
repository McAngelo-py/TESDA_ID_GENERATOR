import os
import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from PIL import Image, ImageTk
from docx import Document
import csv
import zipfile
import xml.etree.ElementTree as ET


# ── Theme definitions ─────────────────────────────────────────────────────────
THEMES = {
    "light": {
        "BG":          "#f8fafc",
        "PANEL":       "#ffffff",
        "BORDER":      "#e2e8f0",
        "ENTRY_BG":    "#ffffff",
        "ENTRY_FG":    "#1e293b",
        "TEXT":        "#0f172a",
        "MUTED":       "#64748b",
        "TOPBAR":      "#ffffff",
        "CARD_TITLE":  "#475569",
        "HINT_FG":     "#94a3b8",
        "toggle_icon": "🌙  Dark Mode",
    },
    "dark": {
        "BG":          "#0f172a",
        "PANEL":       "#1e293b",
        "BORDER":      "#334155",
        "ENTRY_BG":    "#0f172a",
        "ENTRY_FG":    "#f1f5f9",
        "TEXT":        "#f8fafc",
        "MUTED":       "#94a3b8",
        "TOPBAR":      "#1e293b",
        "CARD_TITLE":  "#cbd5e1",
        "HINT_FG":     "#475569",
        "toggle_icon": "☀  Light Mode",
    },
}

ACCENT  = "#3b82f6"  # Modern Blue
ACCENT2 = "#8b5cf6"  # Modern Violet
TEAL    = "#14b8a6"  # Modern Teal
SUCCESS = "#10b981"  # Modern Emerald
DANGER  = "#ef4444"  # Modern Red
WARNING = "#f59e0b"  # Modern Amber

FONT_SUB    = ("Segoe UI", 10, "bold")
FONT_BODY   = ("Segoe UI", 10)
FONT_MONO   = ("Consolas", 10)
FONT_MONO_S = ("Consolas", 9)
FONT_LABEL  = ("Segoe UI", 9, "bold")


def _lighten(hex_color, amount=28):
    h = hex_color.lstrip("#")
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    r, g, b = min(r + amount, 255), min(g + amount, 255), min(b + amount, 255)
    return f"#{r:02x}{g:02x}{b:02x}"


class _TextStub:
    """Drop-in stub for a removed tk.Text widget."""
    def __init__(self): self._data = ""
    def get(self, *a, **kw): return self._data
    def insert(self, idx, text): self._data += text
    def delete(self, *a, **kw): self._data = ""


class IDGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("ID Generator")
        self.root.geometry("1200x840")
        self.root.minsize(900, 600)
        self.root.state('zoomed')
        self.root.resizable(True, True)

        self._theme_name = "light"
        self._T = THEMES["light"]
        self._themed_widgets = []
        self._themed_buttons = []

        self.template_path = tk.StringVar()
        self.name_placeholder_var = tk.StringVar(value="NAME HERE")
        self.course_var = tk.StringVar(value="COURSE HERE")
        self.detected_name_placeholder_count = 0
        self.detected_id_placeholder_count = 0
        self.detected_address_placeholder_count = 0
        self.detected_blood_placeholder_count = 0
        self.detected_sex_placeholder_count = 0
        self.detected_gender_placeholder_count = 0
        self.detected_emergency_name_placeholder_count = 0
        self.detected_emergency_number_placeholder_count = 0
        self.detected_emergency_address_placeholder_count = 0
        self.detected_course_placeholder_count = 0
        self.id_seed_value = "2026-000"
        self.repeat_factor = 1
        self.gender_text = _TextStub()  # gender input removed from UI

        self.setup_ui()

    # ── Theme engine ──────────────────────────────────────────────────────────

    def _reg(self, widget, role):
        self._themed_widgets.append((widget, role))
        return widget

    def _reg_btn(self, btn, color):
        self._themed_buttons.append((btn, color))
        return btn

    def _apply_theme(self):
        T = self._T
        self.root.configure(bg=T["BG"])
        self._canvas.configure(bg=T["BG"])
        for widget, role in self._themed_widgets:
            try:
                if role == "bg":
                    widget.configure(bg=T["BG"])
                elif role == "panel":
                    widget.configure(bg=T["PANEL"])
                elif role == "border":
                    widget.configure(bg=T["BORDER"])
                elif role == "topbar":
                    widget.configure(bg=T["TOPBAR"])
                elif role == "text_topbar":
                    widget.configure(bg=T["TOPBAR"], fg=T["TEXT"])
                elif role == "muted_topbar":
                    widget.configure(bg=T["TOPBAR"], fg=T["MUTED"])
                elif role == "muted":
                    widget.configure(bg=T["PANEL"], fg=T["MUTED"])
                elif role == "card_title":
                    widget.configure(bg=T["PANEL"], fg=T["CARD_TITLE"])
                elif role == "hint":
                    widget.configure(bg=T["PANEL"], fg=T["HINT_FG"])
                elif role == "status":
                    widget.configure(bg=T["PANEL"])
                elif role == "entry":
                    widget.configure(
                        bg=T["ENTRY_BG"], fg=T["ENTRY_FG"],
                        insertbackground=T["ENTRY_FG"],
                        highlightbackground=T["BORDER"],
                    )
                elif role == "scrolled":
                    widget.configure(
                        bg=T["ENTRY_BG"], fg=T["ENTRY_FG"],
                        insertbackground=T["ENTRY_FG"],
                        highlightbackground=T["BORDER"],
                    )
            except tk.TclError:
                pass
        for btn, color in self._themed_buttons:
            try:
                btn.configure(bg=color, activebackground=color)
            except tk.TclError:
                pass
        self.toggle_btn.configure(
            text=T["toggle_icon"],
            bg=T["TOPBAR"], fg=T["MUTED"],
            activebackground=T["TOPBAR"], activeforeground=ACCENT,
        )

    def toggle_theme(self):
        self._theme_name = "light" if self._theme_name == "dark" else "dark"
        self._T = THEMES[self._theme_name]
        self._apply_theme()

    # ── Widget factories ──────────────────────────────────────────────────────

    def _styled_button(self, parent, text, command, color=ACCENT):
        btn = tk.Button(
            parent, text=text, command=command,
            bg=color, fg="#ffffff",
            activebackground=_lighten(color), activeforeground="#ffffff",
            font=FONT_SUB, padx=16, pady=8,
            relief="flat", bd=0, cursor="hand2",
        )
        def _on(e):  btn.config(bg=_lighten(color))
        def _off(e): btn.config(bg=color)
        btn.bind("<Enter>", _on)
        btn.bind("<Leave>", _off)
        self._reg_btn(btn, color)
        return btn

    def _make_entry(self, parent, textvariable=None, width=22):
        T = self._T
        e = tk.Entry(
            parent, textvariable=textvariable,
            bg=T["ENTRY_BG"], fg=T["ENTRY_FG"],
            insertbackground=T["ENTRY_FG"],
            relief="flat", bd=0, font=FONT_MONO,
            width=width, highlightthickness=1,
            highlightcolor=ACCENT, highlightbackground=T["BORDER"],
        )
        self._reg(e, "entry")
        return e

    def _make_dropdown(self, parent, variable, choices):
        T = self._T
        om = tk.OptionMenu(parent, variable, *choices)
        om.config(
            bg=T["ENTRY_BG"], fg=T["ENTRY_FG"],
            activebackground=T["PANEL"], activeforeground=T["TEXT"],
            highlightthickness=1, highlightbackground=T["BORDER"],
            relief="flat", bd=0, font=FONT_MONO, cursor="hand2",
            indicatoron=True,
            padx=10, pady=4
        )
        om["menu"].config(
            bg=T["ENTRY_BG"], fg=T["ENTRY_FG"],
            activebackground=ACCENT, activeforeground="#fff",
            font=FONT_MONO, relief="flat", bd=0,
        )
        self._themed_buttons.append((om, T["ENTRY_BG"]))
        return om

    def _make_scrolled(self, parent, height=8, width=20):
        T = self._T
        box = scrolledtext.ScrolledText(
            parent, height=height, width=width, font=FONT_MONO,
            bg=T["ENTRY_BG"], fg=T["ENTRY_FG"],
            insertbackground=T["ENTRY_FG"],
            relief="flat", bd=0,
            highlightthickness=1, highlightcolor=ACCENT,
            highlightbackground=T["BORDER"],
            selectbackground=ACCENT, selectforeground="#fff",
            padx=10, pady=10
        )
        self._reg(box, "scrolled")
        return box

    def _card(self, parent, title=""):
        T = self._T
        outer = tk.Frame(parent, bg=T["BORDER"], bd=0)
        self._reg(outer, "border")
        inner = tk.Frame(outer, bg=T["PANEL"], padx=20, pady=18)
        self._reg(inner, "panel")
        inner.pack(fill="both", expand=True, padx=1, pady=1)
        if title:
            lbl = tk.Label(inner, text=title, bg=T["PANEL"], fg=T["CARD_TITLE"],
                           font=FONT_LABEL, anchor="w")
            lbl.pack(fill="x", pady=(0, 14))
            self._reg(lbl, "card_title")
        return outer, inner

    def _col_label(self, parent, text):
        T = self._T
        lbl = tk.Label(parent, text=text, bg=T["PANEL"], fg=T["MUTED"], font=FONT_LABEL)
        self._reg(lbl, "muted")
        return lbl

    # ── UI Construction ───────────────────────────────────────────────────────

    def setup_ui(self):
        T = self._T

        # Topbar
        topbar = tk.Frame(self.root, bg=T["TOPBAR"], height=70)
        self._reg(topbar, "topbar")
        topbar.pack(fill="x", side="top")
        topbar.pack_propagate(False)

        tk.Frame(topbar, bg=ACCENT, width=5).pack(side="left", fill="y")

        # Logo
        try:
            logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
            img = Image.open(logo_path)
            img = img.resize((45, 45), Image.Resampling.LANCZOS)
            self.logo_img = ImageTk.PhotoImage(img)
            logo_lbl = tk.Label(topbar, image=self.logo_img, bg=T["TOPBAR"])
            logo_lbl.pack(side="left", padx=(18, 0))
        except Exception as e:
            print(f"Logo error: {e}")

        title_lbl = tk.Label(topbar, text="TESDA ID Generator",
                             bg=T["TOPBAR"], fg=T["TEXT"],
                             font=("Segoe UI", 16, "bold"), padx=12)
        title_lbl.pack(side="left", pady=10)
        self._reg(title_lbl, "text_topbar")

        sub_lbl = tk.Label(topbar, text="Batch replacement for .docx templates",
                           bg=T["TOPBAR"], fg=T["MUTED"], font=FONT_BODY)
        sub_lbl.pack(side="left")
        self._reg(sub_lbl, "muted_topbar")

        self.toggle_btn = tk.Button(
            topbar, text=T["toggle_icon"], command=self.toggle_theme,
            bg=T["TOPBAR"], fg=T["MUTED"],
            activebackground=T["TOPBAR"], activeforeground=ACCENT,
            font=FONT_BODY, padx=16, pady=8,
            relief="flat", bd=0, cursor="hand2",
        )
        self.toggle_btn.pack(side="right", padx=20)

        sep = tk.Frame(self.root, bg=T["BORDER"], height=1)
        self._reg(sep, "border")
        sep.pack(fill="x")

        # Scrollable canvas
        canvas_container = tk.Frame(self.root, bg=T["BG"])
        self._reg(canvas_container, "bg")
        canvas_container.pack(fill="both", expand=True)

        self._canvas = tk.Canvas(canvas_container, bg=T["BG"], bd=0, highlightthickness=0)
        vscroll = tk.Scrollbar(canvas_container, orient="vertical", command=self._canvas.yview)
        self._canvas.configure(yscrollcommand=vscroll.set)
        vscroll.pack(side="right", fill="y")
        self._canvas.pack(side="left", fill="both", expand=True)

        _body_frame = tk.Frame(self._canvas, bg=T["BG"])
        self._reg(_body_frame, "bg")
        self._body_window = self._canvas.create_window((0, 0), window=_body_frame, anchor="nw")

        def _on_frame_configure(e):
            self._canvas.configure(scrollregion=self._canvas.bbox("all"))
        def _on_canvas_resize(e):
            self._canvas.itemconfig(self._body_window, width=e.width)
        _body_frame.bind("<Configure>", _on_frame_configure)
        self._canvas.bind("<Configure>", _on_canvas_resize)

        def _on_mousewheel(e):
            self._canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
        def _on_mousewheel_linux(e):
            self._canvas.yview_scroll(-1 if e.num == 4 else 1, "units")
        self._canvas.bind_all("<MouseWheel>", _on_mousewheel)
        self._canvas.bind_all("<Button-4>", _on_mousewheel_linux)
        self._canvas.bind_all("<Button-5>", _on_mousewheel_linux)

        body = tk.Frame(_body_frame, bg=T["BG"])
        self._reg(body, "bg")
        body.pack(fill="both", expand=True, padx=32, pady=28)

        # Card 01: Template
        c1_out, c1_in = self._card(body, "01  TEMPLATE")
        c1_out.pack(fill="x", pady=(0, 20))

        row = tk.Frame(c1_in, bg=T["PANEL"])
        self._reg(row, "panel")
        row.pack(fill="x")

        self.upload_btn = self._styled_button(row, "⊕  Upload .docx", self.upload_file, ACCENT)
        self.upload_btn.pack(side="left")

        self.file_label = tk.Label(row, text="No file selected",
                                   bg=T["PANEL"], fg=T["MUTED"],
                                   font=FONT_MONO_S, padx=16)
        self._reg(self.file_label, "muted")
        self.file_label.pack(side="left")

        ph_lbl = tk.Label(row, text="Name placeholder:", bg=T["PANEL"], fg=T["MUTED"], font=FONT_BODY)
        self._reg(ph_lbl, "muted")
        ph_lbl.pack(side="right", padx=(0, 8))
        ph_entry = self._make_entry(row, textvariable=self.name_placeholder_var, width=22)
        ph_entry.pack(side="right")

        # Card 02: Course selector
        COURSES = ['Agroentrepreneurship NC II', 'Agroentrepreneurship NC III', 'Barangay Health Services NC II', 'Bookkeeping NC III', "Community-Based Trainer's Methodology Course", 'Dressmaking NC II', 'Driving NC II', 'Early Childhood Care and Development Services NC III', 'Electrical Installation and Maintenance NC II', 'Electrical Installation and Maintenance NC III', 'Housekeeping NC II', 'Organic Agriculture Production NC II', 'PV Systems Installation NC II', 'Shielded Metal Arc Welding NC I', 'Shielded Metal Arc Welding NC II', "Trainer's Methodology Level I"]
        cc_out, cc_in = self._card(body, "02  COURSE")
        cc_out.pack(fill="x", pady=(0, 20))

        course_row = tk.Frame(cc_in, bg=T["PANEL"])
        self._reg(course_row, "panel")
        course_row.pack(fill="x")

        c_lbl = tk.Label(course_row, text="Select course:", bg=T["PANEL"], fg=T["MUTED"], font=FONT_BODY)
        self._reg(c_lbl, "muted")
        c_lbl.pack(side="left", padx=(0, 12))

        self.course_dropdown = self._make_dropdown(course_row, self.course_var, COURSES)
        self.course_dropdown.pack(side="left", fill="x", expand=True)

        # Card 03: Detected
        c2_out, c2_in = self._card(body, "03  DETECTED PLACEHOLDERS")
        c2_out.pack(fill="x", pady=(0, 20))

        self.detected_names_text = self._make_scrolled(c2_in, height=5)
        self.detected_names_text.config(state="disabled")
        self.detected_names_text.pack(fill="x")

        btn_row = tk.Frame(c2_in, bg=T["PANEL"])
        self._reg(btn_row, "panel")
        btn_row.pack(fill="x", pady=(14, 0))

        self.extract_btn = self._styled_button(btn_row, "⟳  Auto-detect Placeholders",
                                               self.autofill_names_from_docx, ACCENT2)
        self.extract_btn.pack(side="left")

        self.csv_btn = self._styled_button(btn_row, "⊕  Upload CSV", self.upload_csv, TEAL)
        self.csv_btn.pack(side="left", padx=(12, 0))

        # Card 04: Person data
        c3_out, c3_in = self._card(body, "04  PERSON DATA  —  one entry per line")
        c3_out.pack(fill="both", expand=True, pady=(0, 20))

        # Row A
        row_a = tk.Frame(c3_in, bg=T["PANEL"])
        self._reg(row_a, "panel")
        row_a.pack(fill="both", expand=True, pady=(0, 12))

        for title, default, attr in [
            ("Name",        "NAME HERE",         "name_text"),
            ("Address",     "HOME ADDRESS HERE",  "address_text"),
            ("Blood Type",  "BLOOD TYPE HERE",    "blood_text"),
            ("Sex",         "SEX HERE",           "sex_text"),
        ]:
            col = tk.Frame(row_a, bg=T["PANEL"])
            self._reg(col, "panel")
            col.pack(side="left", fill="both", expand=True, padx=(0, 10))
            self._col_label(col, title).pack(anchor="w", pady=(0, 6))
            box = self._make_scrolled(col, height=7, width=14)
            box.insert(tk.END, default)
            box.pack(fill="both", expand=True)
            setattr(self, attr, box)

        div = tk.Frame(c3_in, bg=T["BORDER"], height=1)
        self._reg(div, "border")
        div.pack(fill="x", pady=(0, 12))

        # Row B
        row_b = tk.Frame(c3_in, bg=T["PANEL"])
        self._reg(row_b, "panel")
        row_b.pack(fill="both", expand=True)

        for title, default, attr in [
            ("Emergency Name",    "EMERGENCY NAME HERE",    "emergency_name_text"),
            ("Emergency Number",  "EMERGENCY NUMBER HERE",  "emergency_number_text"),
            ("Emergency Address", "EMERGENCY ADDRESS HERE", "emergency_address_text"),
        ]:
            col = tk.Frame(row_b, bg=T["PANEL"])
            self._reg(col, "panel")
            col.pack(side="left", fill="both", expand=True, padx=(0, 10))
            self._col_label(col, title).pack(anchor="w", pady=(0, 6))
            box = self._make_scrolled(col, height=7, width=14)
            box.insert(tk.END, default)
            box.pack(fill="both", expand=True)
            setattr(self, attr, box)

        hint = tk.Label(c3_in, text="ℹ  IDs are auto-generated (+1) starting from  2026-000",
                        bg=T["PANEL"], fg=T["HINT_FG"], font=FONT_BODY, anchor="w")
        self._reg(hint, "hint")
        hint.pack(fill="x", pady=(14, 0))

        # Card 05: Generate
        c4_out, c4_in = self._card(body, "05  GENERATE")
        c4_out.pack(fill="x")

        self.generate_btn = self._styled_button(c4_in, "⚡  Generate Updated File",
                                                self.process_files, DANGER)
        self.generate_btn.config(font=("Segoe UI", 12, "bold"), pady=14)
        self.generate_btn.pack(fill="x")

        self.status_label = tk.Label(c4_in, text="",
                                     bg=T["PANEL"], fg=T["MUTED"], font=FONT_BODY)
        self._reg(self.status_label, "status")
        self.status_label.pack(pady=(12, 2))

    # ── All logic methods below are UNCHANGED ─────────────────────────────────

    def upload_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if file_path:
            self.template_path.set(file_path)
            self.file_label.config(text=os.path.basename(file_path),
                                   font=FONT_MONO_S, fg=self._T["TEXT"])

    def upload_csv(self):
        file_path = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
        if file_path:
            try:
                with open(file_path, mode='r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    names = []
                    addresses = []
                    bloods = []
                    sexes = []
                    genders = []
                    emergency_names = []
                    emergency_numbers = []
                    emergency_addresses = []
                    for row in reader:
                        if not row:
                            continue
                        if len(row) >= 8:
                            name = row[0].strip()
                            address = row[1].strip()
                            blood = row[2].strip()
                            sex = row[3].strip()
                            gender = row[4].strip()
                            emergency_name = row[5].strip()
                            emergency_number = row[6].strip()
                            emergency_address = row[7].strip()
                            if name:
                                names.append(name)
                                addresses.append(address)
                                bloods.append(blood)
                                sexes.append(sex)
                                genders.append(gender)
                                emergency_names.append(emergency_name)
                                emergency_numbers.append(emergency_number)
                                emergency_addresses.append(emergency_address)
                        elif len(row) >= 4:
                            name = row[0].strip()
                            address = row[1].strip()
                            blood = row[2].strip()
                            sex = row[3].strip()
                            if name:
                                names.append(name)
                                addresses.append(address)
                                bloods.append(blood)
                                sexes.append(sex)
                                genders.append("")
                                emergency_names.append("")
                                emergency_numbers.append("")
                                emergency_addresses.append("")
                        elif len(row) >= 1 and row[0].strip():
                            names.append(row[0].strip())
                            addresses.append("")
                            bloods.append("")
                            sexes.append("")
                            genders.append("")
                            emergency_names.append("")
                            emergency_numbers.append("")
                            emergency_addresses.append("")
                    if names:
                        self.name_text.delete("1.0", tk.END)
                        self.address_text.delete("1.0", tk.END)
                        self.blood_text.delete("1.0", tk.END)
                        self.sex_text.delete("1.0", tk.END)
                        self.gender_text.delete("1.0", tk.END)
                        self.emergency_name_text.delete("1.0", tk.END)
                        self.emergency_number_text.delete("1.0", tk.END)
                        self.emergency_address_text.delete("1.0", tk.END)
                        self.name_text.insert(tk.END, "\n".join(names))
                        self.address_text.insert(tk.END, "\n".join(addresses))
                        self.blood_text.insert(tk.END, "\n".join(bloods))
                        self.sex_text.insert(tk.END, "\n".join(sexes))
                        self.gender_text.insert(tk.END, "\n".join(genders))
                        self.emergency_name_text.insert(tk.END, "\n".join(emergency_names))
                        self.emergency_number_text.insert(tk.END, "\n".join(emergency_numbers))
                        self.emergency_address_text.insert(tk.END, "\n".join(emergency_addresses))
                        messagebox.showinfo("Success", f"Loaded {len(names)} person row(s) from CSV.")
            except Exception as e:
                messagebox.showerror("Error", f"Could not read CSV: {str(e)}")

    def parse_replacement_rows(self):
        if self.detected_name_placeholder_count <= 0:
            raise ValueError("Click 'Auto-detect Placeholders' first to detect name placeholders.")
        name_lines = [line.strip() for line in self.name_text.get("1.0", tk.END).split("\n")]
        address_lines = [line.strip() for line in self.address_text.get("1.0", tk.END).split("\n")]
        blood_lines = [line.strip() for line in self.blood_text.get("1.0", tk.END).split("\n")]
        sex_lines = [line.strip() for line in self.sex_text.get("1.0", tk.END).split("\n")]
        gender_lines = [line.strip() for line in self.gender_text.get("1.0", tk.END).split("\n")]
        emergency_name_lines = [line.strip() for line in self.emergency_name_text.get("1.0", tk.END).split("\n")]
        emergency_number_lines = [line.strip() for line in self.emergency_number_text.get("1.0", tk.END).split("\n")]
        emergency_address_lines = [line.strip() for line in self.emergency_address_text.get("1.0", tk.END).split("\n")]
        clean_names = [line for line in name_lines if line and line.upper() != "NAME HERE"]
        if not clean_names:
            raise ValueError("Please input names (one per line).")
        clean_addresses = [line for line in address_lines if line.upper() != "HOME ADDRESS HERE"]
        clean_bloods = [line for line in blood_lines if line.upper() != "BLOOD TYPE HERE"]
        clean_sexes = [line for line in sex_lines if line.upper() != "SEX HERE"]
        clean_genders = [line for line in gender_lines if line.upper() != "GENDER HERE"]
        clean_emergency_names = [line for line in emergency_name_lines if line.upper() != "EMERGENCY NAME HERE"]
        clean_emergency_numbers = [line for line in emergency_number_lines if line.upper() != "EMERGENCY NUMBER HERE"]
        clean_emergency_addresses = [line for line in emergency_address_lines if line.upper() != "EMERGENCY ADDRESS HERE"]
        expected_people = max(1, self.detected_name_placeholder_count // self.repeat_factor)
        if len(clean_names) > expected_people:
            placeholder = self.name_placeholder_var.get().strip() or "NAME HERE"
            raise ValueError(
                f"Too many names: {len(clean_names)} provided, but only "
                f"{expected_people} person entry(ies) expected from {self.detected_name_placeholder_count} '{placeholder}' placeholders."
            )
        expanded_names = []
        expanded_addresses = []
        expanded_blood_types = []
        expanded_sexes = []
        expanded_genders = []
        expanded_emergency_names = []
        expanded_emergency_numbers = []
        expanded_emergency_addresses = []
        for idx, name in enumerate(clean_names):
            address = clean_addresses[idx] if idx < len(clean_addresses) else ""
            blood = clean_bloods[idx] if idx < len(clean_bloods) else ""
            sex = clean_sexes[idx] if idx < len(clean_sexes) else ""
            gender = clean_genders[idx] if idx < len(clean_genders) else ""
            emergency_name = clean_emergency_names[idx] if idx < len(clean_emergency_names) else ""
            emergency_number = clean_emergency_numbers[idx] if idx < len(clean_emergency_numbers) else ""
            emergency_address = clean_emergency_addresses[idx] if idx < len(clean_emergency_addresses) else ""
            for _ in range(self.repeat_factor):
                expanded_names.append(name)
                expanded_addresses.append(address)
                expanded_blood_types.append(blood)
                expanded_sexes.append(sex)
                expanded_genders.append(gender)
                expanded_emergency_names.append(emergency_name)
                expanded_emergency_numbers.append(emergency_number)
                expanded_emergency_addresses.append(emergency_address)
        skipped = expected_people - len(clean_names)
        return (
            expanded_names, expanded_addresses, expanded_blood_types,
            expanded_sexes, expanded_genders, expanded_emergency_names,
            expanded_emergency_numbers, expanded_emergency_addresses,
            skipped, len(clean_names), expected_people,
        )

    def build_auto_ids(self, count, seed_value="2026-000"):
        if count <= 0:
            return []
        if "-" in seed_value:
            prefix, numeric = seed_value.rsplit("-", 1)
            if numeric.isdigit():
                width = len(numeric)
                start = int(numeric)
                return [f"{prefix}-{str(start + i).zfill(width)}" for i in range(count)]
        return [seed_value for _ in range(count)]

    def autofill_names_from_docx(self):
        template = self.template_path.get()
        if not template:
            messagebox.showerror("Error", "Please select a .docx template file first.")
            return
        try:
            name_placeholder = self.name_placeholder_var.get().strip() or "NAME HERE"
            name_placeholder_count = self.count_placeholders_in_docx(template, name_placeholder)
            id_placeholder_count = self.count_placeholders_in_docx(template, "2026-000")
            address_placeholder_count = self.count_placeholders_in_docx(template, "HOME ADDRESS HERE") + self.count_placeholders_in_docx(template, "HOME ADRESS HERE")
            blood_placeholder_count = self.count_placeholders_in_docx(template, "BLOOD TYPE HERE")
            sex_placeholder_count = self.count_placeholders_in_docx(template, "SEX HERE")
            gender_placeholder_count = self.count_placeholders_in_docx(template, "GENDER HERE")
            emergency_name_placeholder_count = self.count_placeholders_in_docx(template, "EMERGENCY NAME HERE")
            emergency_number_placeholder_count = self.count_placeholders_in_docx(template, "EMERGENCY NUMBER HERE")
            emergency_address_placeholder_count = self.count_placeholders_in_docx(template, "EMERGENCY ADDRESS HERE")
            course_placeholder_count = self.count_placeholders_in_docx(template, "COURSE HERE")
            if name_placeholder_count <= 0:
                messagebox.showwarning("No Placeholder Found", f"No '{name_placeholder}' placeholder was found in the selected .docx.")
                return
            self.detected_name_placeholder_count = name_placeholder_count
            self.detected_id_placeholder_count = id_placeholder_count
            self.detected_address_placeholder_count = address_placeholder_count
            self.detected_blood_placeholder_count = blood_placeholder_count
            self.detected_sex_placeholder_count = sex_placeholder_count
            self.detected_gender_placeholder_count = gender_placeholder_count
            self.detected_emergency_name_placeholder_count = emergency_name_placeholder_count
            self.detected_emergency_number_placeholder_count = emergency_number_placeholder_count
            self.detected_emergency_address_placeholder_count = emergency_address_placeholder_count
            self.detected_course_placeholder_count = course_placeholder_count
            self.repeat_factor = 2 if name_placeholder_count % 2 == 0 and id_placeholder_count % 2 == 0 else 1
            expected_people = max(1, name_placeholder_count // self.repeat_factor)
            self.detected_names_text.config(state="normal")
            self.detected_names_text.delete("1.0", tk.END)
            self.detected_names_text.insert(
                tk.END,
                f"{name_placeholder} x {name_placeholder_count}\n"
                f"2026-000 x {id_placeholder_count}\n"
                f"HOME ADDRESS HERE x {address_placeholder_count}\n"
                f"BLOOD TYPE HERE x {blood_placeholder_count}\n"
                f"SEX HERE x {sex_placeholder_count}\n"
                f"GENDER HERE x {gender_placeholder_count}\n"
                f"EMERGENCY NAME HERE x {emergency_name_placeholder_count}\n"
                f"EMERGENCY NUMBER HERE x {emergency_number_placeholder_count}\n"
                f"EMERGENCY ADDRESS HERE x {emergency_address_placeholder_count}\n"
                f"COURSE HERE x {course_placeholder_count}\n"
                f"Input entries expected: {expected_people} (repeat factor: {self.repeat_factor}x)"
            )
            self.detected_names_text.config(state="disabled")
            self.name_text.delete("1.0", tk.END)
            self.address_text.delete("1.0", tk.END)
            self.blood_text.delete("1.0", tk.END)
            self.sex_text.delete("1.0", tk.END)
            self.gender_text.delete("1.0", tk.END)
            self.emergency_name_text.delete("1.0", tk.END)
            self.emergency_number_text.delete("1.0", tk.END)
            self.emergency_address_text.delete("1.0", tk.END)
            self.name_text.insert(tk.END, "\n".join([name_placeholder for _ in range(expected_people)]))
            self.address_text.insert(tk.END, "\n".join(["HOME ADDRESS HERE" for _ in range(expected_people)]))
            self.blood_text.insert(tk.END, "\n".join(["BLOOD TYPE HERE" for _ in range(expected_people)]))
            self.sex_text.insert(tk.END, "\n".join(["SEX HERE" for _ in range(expected_people)]))
            self.gender_text.insert(tk.END, "\n".join(["GENDER HERE" for _ in range(expected_people)]))
            self.emergency_name_text.insert(tk.END, "\n".join(["EMERGENCY NAME HERE" for _ in range(expected_people)]))
            self.emergency_number_text.insert(tk.END, "\n".join(["EMERGENCY NUMBER HERE" for _ in range(expected_people)]))
            self.emergency_address_text.insert(tk.END, "\n".join(["EMERGENCY ADDRESS HERE" for _ in range(expected_people)]))
            messagebox.showinfo(
                "Placeholder Loaded",
                f"Detected placeholders for names, IDs, address, blood type, sex, gender, and emergency fields.\n"
                f"Using repeat factor {self.repeat_factor}x, so type {expected_people} person row(s) only. IDs auto-increment by +1 per person."
            )
        except Exception as e:
            messagebox.showerror("Error", f"Could not extract names from .docx: {str(e)}")

    def count_placeholders_in_docx(self, source_path, placeholder):
        count = 0
        placeholder_upper = placeholder.upper()
        placeholder_words = placeholder_upper.split()
        try:
            with zipfile.ZipFile(source_path, "r") as archive:
                xml_parts = [name for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml")]
                namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                for part_name in xml_parts:
                    root = ET.fromstring(archive.read(part_name))
                    text_nodes = [node for node in root.findall(".//w:t", namespace) if node.text]
                    i = 0
                    while i < len(text_nodes):
                        current = " ".join(text_nodes[i].text.split()).upper()
                        if current == placeholder_upper:
                            count += 1
                            i += 1
                            continue
                        if len(placeholder_words) > 1 and i + len(placeholder_words) - 1 < len(text_nodes):
                            if all(" ".join(text_nodes[i + j].text.split()).upper() == placeholder_words[j] for j in range(len(placeholder_words))):
                                count += 1
                                i += len(placeholder_words)
                                continue
                        i += 1
        except Exception:
            doc = Document(source_path)
            for p in doc.paragraphs:
                if " ".join(p.text.split()).upper() == placeholder_upper:
                    count += 1
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if " ".join(p.text.split()).upper() == placeholder_upper:
                                count += 1
        return count

    def replace_placeholders_in_docx_xml(
        self, source_path, output_path,
        replacement_names, replacement_ids, replacement_addresses,
        replacement_blood_types, replacement_sexes, replacement_genders,
        replacement_emergency_names, replacement_emergency_numbers,
        replacement_emergency_addresses, name_placeholder,
        replacement_courses=None,
    ):
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        text_tag = f"{{{w_ns}}}t"
        name_idx = id_idx = address_idx = blood_idx = 0
        sex_idx = gender_idx = emergency_name_idx = 0
        emergency_number_idx = emergency_address_idx = course_idx = 0
        if replacement_courses is None:
            replacement_courses = []
        name_placeholder_upper = name_placeholder.upper()
        name_placeholder_words = name_placeholder_upper.split()
        with zipfile.ZipFile(source_path, "r") as source_zip:
            with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as out_zip:
                for item in source_zip.infolist():
                    data = source_zip.read(item.filename)
                    if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                        try:
                            root = ET.fromstring(data)
                            changed = False
                            text_nodes = [node for node in root.iter(text_tag) if node.text]
                            i = 0
                            while i < len(text_nodes):
                                normalized = " ".join(text_nodes[i].text.split()).upper()
                                if normalized == name_placeholder_upper and name_idx < len(replacement_names):
                                    text_nodes[i].text = replacement_names[name_idx]
                                    name_idx += 1; changed = True; i += 1; continue
                                if len(name_placeholder_words) > 1 and i + len(name_placeholder_words) - 1 < len(text_nodes) and name_idx < len(replacement_names):
                                    if all(" ".join(text_nodes[i + j].text.split()).upper() == name_placeholder_words[j] for j in range(len(name_placeholder_words))):
                                        text_nodes[i].text = replacement_names[name_idx]
                                        for j in range(1, len(name_placeholder_words)):
                                            text_nodes[i + j].text = ""
                                        name_idx += 1; changed = True; i += len(name_placeholder_words); continue
                                if normalized == "2026-000" and id_idx < len(replacement_ids):
                                    text_nodes[i].text = replacement_ids[id_idx]; id_idx += 1; changed = True
                                elif normalized in ("HOME ADDRESS HERE", "HOME ADRESS HERE") and address_idx < len(replacement_addresses):
                                    text_nodes[i].text = replacement_addresses[address_idx]; address_idx += 1; changed = True
                                elif normalized == "BLOOD TYPE HERE" and blood_idx < len(replacement_blood_types):
                                    text_nodes[i].text = replacement_blood_types[blood_idx]; blood_idx += 1; changed = True
                                elif normalized == "SEX HERE" and sex_idx < len(replacement_sexes):
                                    text_nodes[i].text = replacement_sexes[sex_idx]; sex_idx += 1; changed = True
                                elif normalized == "GENDER HERE" and gender_idx < len(replacement_genders):
                                    text_nodes[i].text = replacement_genders[gender_idx]; gender_idx += 1; changed = True
                                elif normalized == "EMERGENCY NAME HERE" and emergency_name_idx < len(replacement_emergency_names):
                                    text_nodes[i].text = replacement_emergency_names[emergency_name_idx]; emergency_name_idx += 1; changed = True
                                elif normalized == "EMERGENCY NUMBER HERE" and emergency_number_idx < len(replacement_emergency_numbers):
                                    text_nodes[i].text = replacement_emergency_numbers[emergency_number_idx]; emergency_number_idx += 1; changed = True
                                elif normalized == "EMERGENCY ADDRESS HERE" and emergency_address_idx < len(replacement_emergency_addresses):
                                    text_nodes[i].text = replacement_emergency_addresses[emergency_address_idx]; emergency_address_idx += 1; changed = True
                                elif normalized == "COURSE HERE" and course_idx < len(replacement_courses):
                                    text_nodes[i].text = replacement_courses[course_idx]; course_idx += 1; changed = True
                                i += 1
                            if changed:
                                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                        except Exception:
                            pass
                    out_zip.writestr(item, data)
        return (name_idx, id_idx, address_idx, blood_idx, sex_idx, gender_idx,
                emergency_name_idx, emergency_number_idx, emergency_address_idx, course_idx)

    def process_files(self):
        template = self.template_path.get()
        if not template:
            messagebox.showerror("Error", "Please select a .docx template file.")
            return
        try:
            (replacement_names, replacement_addresses, replacement_blood_types,
             replacement_sexes, replacement_genders, replacement_emergency_names,
             replacement_emergency_numbers, replacement_emergency_addresses,
             skipped, replaced_count, total_old) = self.parse_replacement_rows()
        except ValueError as e:
            messagebox.showerror("Invalid Input", str(e)); return
        if not replacement_names:
            messagebox.showwarning("Warning", "No replacement names provided."); return
        person_count = max(1, len(replacement_names) // self.repeat_factor) if replacement_names else 0
        base_ids = self.build_auto_ids(person_count, self.id_seed_value)
        replacement_ids = []
        for value in base_ids:
            replacement_ids.extend([value] * self.repeat_factor)
        replacement_ids = replacement_ids[:self.detected_id_placeholder_count]
        auto_id_count = len(replacement_ids)
        id_skipped = self.detected_id_placeholder_count - auto_id_count
        total_ids = self.detected_id_placeholder_count
        course_value = self.course_var.get()
        replacement_courses = [course_value] * self.detected_course_placeholder_count
        self.status_label.config(text="⏳  Processing… please wait.", fg=WARNING)
        self.root.update_idletasks()
        try:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(os.path.dirname(template), f"UPDATED_IDS_{timestamp}.docx")
            (actual_name_replaced, actual_id_replaced, actual_address_replaced,
             actual_blood_replaced, actual_sex_replaced, actual_gender_replaced,
             actual_emergency_name_replaced, actual_emergency_number_replaced,
             actual_emergency_address_replaced, actual_course_replaced) = self.replace_placeholders_in_docx_xml(
                template, output_path, replacement_names, replacement_ids,
                replacement_addresses, replacement_blood_types, replacement_sexes,
                replacement_genders, replacement_emergency_names,
                replacement_emergency_numbers, replacement_emergency_addresses,
                self.name_placeholder_var.get().strip() or "NAME HERE",
                replacement_courses=replacement_courses,
            )
            out_fname = os.path.basename(output_path)
            self.status_label.config(text=f"✓  Saved as {out_fname}", fg=SUCCESS)
            details = f"Replaced {replaced_count} of {total_old} person entry(ies) ({actual_name_replaced} placeholders)."
            skip_note = f"\nSkipped {skipped} trailing person entry(ies) with no replacement input." if skipped else ""
            id_details = ""
            if total_ids > 0:
                id_details = f"\nAuto-generated IDs from {self.id_seed_value}: replaced {actual_id_replaced} of {total_ids} ID placeholder(s)."
                if id_skipped:
                    id_details += f"\nSkipped {id_skipped} trailing ID(s) with no replacement input."
            other_details = (
                f"\nAddress placeholders replaced: {actual_address_replaced} of {self.detected_address_placeholder_count}"
                f"\nBlood placeholders replaced: {actual_blood_replaced} of {self.detected_blood_placeholder_count}"
                f"\nSex placeholders replaced: {actual_sex_replaced} of {self.detected_sex_placeholder_count}"
                f"\nGender placeholders replaced: {actual_gender_replaced} of {self.detected_gender_placeholder_count}"
                f"\nEmergency name placeholders replaced: {actual_emergency_name_replaced} of {self.detected_emergency_name_placeholder_count}"
                f"\nEmergency number placeholders replaced: {actual_emergency_number_replaced} of {self.detected_emergency_number_placeholder_count}"
                f"\nEmergency address placeholders replaced: {actual_emergency_address_replaced} of {self.detected_emergency_address_placeholder_count}"
            )
            messagebox.showinfo("Success", f"Updated file generated successfully!\nSaved to: {output_path}\n{details}{id_details}{other_details}{skip_note}")
        except Exception as e:
            self.status_label.config(text="✗  Error during processing.", fg=DANGER)
            messagebox.showerror("Error", f"An error occurred: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = IDGeneratorApp(root)
    root.mainloop()